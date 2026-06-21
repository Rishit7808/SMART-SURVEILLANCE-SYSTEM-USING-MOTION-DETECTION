from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# --- Page Setup (A4) ---
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

def add_heading_custom(text, size, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def page_break():
    doc.add_page_break()

# ============ 1. COVER PAGE ============
for _ in range(4):
    doc.add_paragraph()
add_heading_custom("VASAVI COLLEGE OF ENGINEERING", 16)
add_heading_custom("(Autonomous)", 14)
add_heading_custom("IBRAHIMBAGH, HYDERABAD - 31", 12)
doc.add_paragraph()
add_heading_custom("DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING", 14)
doc.add_paragraph()
doc.add_paragraph()
add_heading_custom("COURSE-BASED PROJECT REPORT", 16)
add_heading_custom("ON", 14)
doc.add_paragraph()
add_heading_custom("SMART SURVEILLANCE SYSTEM", 18)
add_heading_custom("USING PYTHON AND OPENCV", 14)
doc.add_paragraph()
doc.add_paragraph()
add_para("Submitted by:", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("[Student Name]  -  [Roll Number]", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para("Under the guidance of", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("[Faculty Name]", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para("Academic Year: 2025-2026", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# ============ 2. CERTIFICATE ============
page_break()
add_heading_custom("CERTIFICATE", 16)
doc.add_paragraph()
add_para("This is to certify that the project report entitled \"Smart Surveillance System using Python and OpenCV\" is a bonafide work carried out by [Student Name], Roll No: [Roll Number], of B.E. Computer Science & Engineering, Vasavi College of Engineering (Autonomous), Hyderabad, in partial fulfillment of the requirements for the Course-Based Project during the academic year 2025-2026.")
doc.add_paragraph()
doc.add_paragraph()
add_para("Project Guide", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("[Faculty Name]", align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("Department of CSE", align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()
add_para("Head of Department", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("[HOD Name]", align=WD_ALIGN_PARAGRAPH.RIGHT)

# ============ 3. DECLARATION ============
page_break()
add_heading_custom("DECLARATION", 16)
doc.add_paragraph()
add_para("I hereby declare that the project report entitled \"Smart Surveillance System using Python and OpenCV\" submitted to the Department of Computer Science & Engineering, Vasavi College of Engineering (Autonomous), Hyderabad, is a record of an original work done by me under the guidance of [Faculty Name], and this project work has not been submitted to any other university or institution for the award of any degree or diploma.")
doc.add_paragraph()
doc.add_paragraph()
add_para("Place: Hyderabad", align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("Date: April 2026", align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()
add_para("[Student Name]", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("[Roll Number]", align=WD_ALIGN_PARAGRAPH.RIGHT)

# ============ 4. ACKNOWLEDGEMENT ============
page_break()
add_heading_custom("ACKNOWLEDGEMENT", 16)
doc.add_paragraph()
add_para("I would like to express my sincere gratitude to my project guide [Faculty Name] for their valuable guidance and constant encouragement throughout this project.")
add_para("I am thankful to the Head of the Department of Computer Science & Engineering for providing the necessary facilities and support.")
add_para("I also extend my thanks to Vasavi College of Engineering for providing an excellent academic environment and resources that made this project possible.")
add_para("Finally, I would like to thank my family and friends for their unwavering support and encouragement.")

# ============ 5. TABLE OF CONTENTS ============
page_break()
add_heading_custom("TABLE OF CONTENTS", 16)
doc.add_paragraph()

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_toc_entry(title, page_num, is_chapter=False, indent_level=0):
    """Add a TOC entry with dot leaders between title and page number."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5

    # Set left indent for sub-sections
    if indent_level > 0:
        p.paragraph_format.left_indent = Cm(1.0)

    # Add right-aligned tab stop with dot leaders
    pPr = p._element.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), '8640')  # ~15.24 cm from left margin in twips
    tabs.append(tab)
    pPr.append(tabs)

    run = p.add_run(f"{title}\t{page_num}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if is_chapter:
        run.bold = True
    return p

# --- Preliminary pages ---
add_toc_entry("Certificate", "ii")
add_toc_entry("Declaration", "iii")
add_toc_entry("Acknowledgement", "iv")
add_toc_entry("Abstract", "vi")
doc.add_paragraph()  # Spacer before chapters

# --- Chapter 1 ---
add_toc_entry("Chapter 1: Introduction", "1", is_chapter=True)
add_toc_entry("1.1  Problem Background", "1", indent_level=1)
add_toc_entry("1.2  Importance of the Problem", "1", indent_level=1)
add_toc_entry("1.3  Real-World Use Case", "2", indent_level=1)
add_toc_entry("1.4  Objectives of the Project", "2", indent_level=1)

# --- Chapter 2 ---
add_toc_entry("Chapter 2: Problem Definition and Requirements", "3", is_chapter=True)
add_toc_entry("2.1  Problem Statement", "3", indent_level=1)
add_toc_entry("2.2  Input and Output Format", "3", indent_level=1)
add_toc_entry("2.3  Constraints and Assumptions", "4", indent_level=1)
add_toc_entry("2.4  Example Input/Output", "4", indent_level=1)

# --- Chapter 3 ---
add_toc_entry("Chapter 3: Algorithm Design", "5", is_chapter=True)
add_toc_entry("3.1  Approach Used", "5", indent_level=1)
add_toc_entry("3.2  Algorithm Steps", "5", indent_level=1)
add_toc_entry("3.3  Pseudocode", "6", indent_level=1)

# --- Chapter 4 ---
add_toc_entry("Chapter 4: Implementation", "7", is_chapter=True)
add_toc_entry("4.1  Programming Language and Tools", "7", indent_level=1)
add_toc_entry("4.2  Code Modules and Functions", "7", indent_level=1)
add_toc_entry("4.3  Key Code Snippets", "8", indent_level=1)

# --- Chapter 5 ---
add_toc_entry("Chapter 5: Results and Analysis", "10", is_chapter=True)
add_toc_entry("5.1  Test Cases", "10", indent_level=1)
add_toc_entry("5.2  Output Analysis", "10", indent_level=1)
add_toc_entry("5.3  Complexity Analysis", "11", indent_level=1)

# --- Chapter 6 ---
add_toc_entry("Chapter 6: Conclusion and Future Work", "12", is_chapter=True)
add_toc_entry("6.1  Conclusion", "12", indent_level=1)
add_toc_entry("6.2  Limitations", "12", indent_level=1)
add_toc_entry("6.3  Future Enhancements", "12", indent_level=1)

doc.add_paragraph()  # Spacer
add_toc_entry("References", "13", is_chapter=True)

# ============ 6. ABSTRACT ============
page_break()
add_heading_custom("ABSTRACT", 16)
doc.add_paragraph()
add_para("Security and surveillance have become critical requirements in modern society. Traditional surveillance systems rely on continuous human monitoring, which is inefficient and prone to errors due to fatigue and distraction. This project presents a Smart Surveillance System developed using Python and the OpenCV library that automates the process of motion detection and alert generation.")
add_para("The system captures real-time video from a webcam and applies computer vision techniques including frame differencing, Gaussian blur, thresholding, and contour detection to identify moving objects. When motion is detected, the system automatically records video clips, captures screenshots, generates audio alerts, and maintains a timestamped log of all events.")
add_para("The system features a real-time dashboard overlay displaying system status, motion count, recording status, and timestamps. It provides keyboard controls for manual recording and screenshot capture. The system is lightweight, runs on standard hardware, and can be deployed for home security, office monitoring, or any environment requiring automated surveillance.")
add_para("Keywords: Smart Surveillance, Motion Detection, OpenCV, Computer Vision, Python, Real-time Monitoring", bold=True)

# ============ CHAPTER 1: INTRODUCTION ============
page_break()
add_heading_custom("CHAPTER 1", 16)
add_heading_custom("INTRODUCTION", 16)
doc.add_paragraph()

add_heading_custom("1.1 Problem Background", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("Surveillance systems have evolved significantly over the past few decades. Traditional CCTV systems require constant human monitoring to detect suspicious activities. With the advancement of computer vision and image processing technologies, it is now possible to build intelligent surveillance systems that can automatically detect motion and generate alerts without human intervention.")
add_para("Python, combined with the OpenCV (Open Source Computer Vision) library, provides a powerful yet accessible platform for developing real-time computer vision applications. OpenCV offers optimized algorithms for image processing, video analysis, and object detection that can run efficiently on standard hardware.")

add_heading_custom("1.2 Importance of the Problem", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("Manual surveillance monitoring suffers from several critical limitations:")
add_bullet("Human operators experience fatigue, reducing detection accuracy over time")
add_bullet("24/7 human monitoring is expensive and resource-intensive")
add_bullet("Delayed response times when threats are not immediately noticed")
add_bullet("No automated record-keeping of security events")
add_para("An automated smart surveillance system addresses these issues by providing continuous, tireless monitoring with instant alert generation and automatic event logging.")

add_heading_custom("1.3 Real-World Use Case", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("Smart surveillance systems have widespread applications:")
add_bullet("Home Security: Detecting intruders when residents are away")
add_bullet("Office Security: Monitoring restricted areas after working hours")
add_bullet("Retail Stores: Detecting shoplifting or unauthorized access")
add_bullet("Warehouse Monitoring: Tracking movement in storage facilities")
add_bullet("Wildlife Monitoring: Detecting animal movement in natural habitats")

add_heading_custom("1.4 Objectives of the Project", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("The primary objectives of this project are:")
add_bullet("To develop a real-time motion detection system using Python and OpenCV")
add_bullet("To implement automatic video recording when motion is detected")
add_bullet("To capture and save screenshots of motion events with timestamps")
add_bullet("To generate audio alerts upon detecting motion")
add_bullet("To maintain a comprehensive log file of all surveillance events")
add_bullet("To provide a user-friendly dashboard overlay on the video feed")

# ============ CHAPTER 2: PROBLEM DEFINITION ============
page_break()
add_heading_custom("CHAPTER 2", 16)
add_heading_custom("PROBLEM DEFINITION AND REQUIREMENTS", 16)
doc.add_paragraph()

add_heading_custom("2.1 Problem Statement", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("Design and implement a cost-effective, automated surveillance system that can detect motion in real-time using a standard webcam, automatically record video evidence, capture screenshots, generate alerts, and maintain event logs - all without requiring constant human monitoring.")

add_heading_custom("2.2 Input and Output Format", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("Input:", bold=True)
add_bullet("Real-time video stream from a webcam (640x480 resolution at 20 FPS)")
add_bullet("User keyboard commands (q: quit, r: toggle recording, s: screenshot)")
add_bullet("Configuration parameters (motion threshold, minimum contour area, cooldown period)")
add_para("Output:", bold=True)
add_bullet("Live video display with motion detection overlay and dashboard")
add_bullet("Motion detection mask visualization")
add_bullet("Recorded video files (.avi format) in surveillance_output/recordings/")
add_bullet("Screenshot images (.jpg format) in surveillance_output/screenshots/")
add_bullet("Event log file (motion_log.txt) with timestamped entries")
add_bullet("Audio alert (beep sound) on motion detection")

add_heading_custom("2.3 Constraints and Assumptions", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("Constraints:", bold=True)
add_bullet("System requires a functional webcam accessible by the operating system")
add_bullet("Minimum contour area of 3000 pixels to filter noise")
add_bullet("3-second cooldown between consecutive alerts to prevent spam")
add_bullet("Motion threshold set to 25 for pixel intensity difference")
add_para("Assumptions:", bold=True)
add_bullet("The camera is stationary (fixed position)")
add_bullet("Lighting conditions remain relatively stable")
add_bullet("The system runs on a Windows operating system")
add_bullet("Python 3.x and OpenCV are properly installed")

add_heading_custom("2.4 Example Input/Output", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("Example Log Output:", bold=True)
p = doc.add_paragraph()
run = p.add_run("[2026-04-27 22:38:58] SYSTEM STARTED: Resolution: 640x480\n[2026-04-27 22:39:05] MOTION DETECTED: Event #1\n[2026-04-27 22:39:05] SCREENSHOT SAVED: screenshots/motion_20260427_223905.jpg\n[2026-04-27 22:39:05] RECORDING STARTED: recordings/recording_20260427_223905.avi\n[2026-04-27 22:39:12] RECORDING STOPPED\n[2026-04-27 22:40:00] SYSTEM STOPPED: Total motions detected: 1")
run.font.name = 'Courier New'
run.font.size = Pt(10)

# ============ CHAPTER 3: ALGORITHM DESIGN ============
page_break()
add_heading_custom("CHAPTER 3", 16)
add_heading_custom("ALGORITHM DESIGN", 16)
doc.add_paragraph()

add_heading_custom("3.1 Approach Used", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("The system uses a Frame Differencing approach for motion detection. This technique compares consecutive frames from the video feed to identify regions where pixel values have changed significantly, indicating movement. The approach involves the following image processing pipeline:")
add_bullet("Frame Differencing: Compute absolute difference between consecutive frames")
add_bullet("Grayscale Conversion: Convert the difference image to grayscale")
add_bullet("Gaussian Blur: Apply blur to reduce noise and minor variations")
add_bullet("Binary Thresholding: Convert to binary image based on intensity threshold")
add_bullet("Dilation: Expand white regions to fill gaps in detected motion areas")
add_bullet("Contour Detection: Find boundaries of motion regions")
add_bullet("Area Filtering: Ignore small contours (noise) below minimum area threshold")

add_heading_custom("3.2 Algorithm Steps", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("Step 1: Initialize the webcam and read the first two frames.")
add_para("Step 2: Compute the absolute difference between Frame1 and Frame2.")
add_para("Step 3: Convert the difference image to grayscale.")
add_para("Step 4: Apply Gaussian Blur (5x5 kernel) to reduce noise.")
add_para("Step 5: Apply binary thresholding (threshold = 25) to isolate motion regions.")
add_para("Step 6: Dilate the binary image (3 iterations) to fill gaps.")
add_para("Step 7: Find contours in the dilated image.")
add_para("Step 8: For each contour with area > 3000 pixels, draw a bounding rectangle and label it.")
add_para("Step 9: If motion is detected and cooldown has elapsed, trigger alerts (sound, screenshot, recording, log entry).")
add_para("Step 10: Display the annotated frame with dashboard overlay.")
add_para("Step 11: Update frames (Frame1 = Frame2, read new Frame2) and repeat from Step 2.")

add_heading_custom("3.3 Pseudocode", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
pseudo = """INITIALIZE webcam
READ frame1, frame2
WHILE webcam is open:
    diff = ABSOLUTE_DIFFERENCE(frame1, frame2)
    gray = CONVERT_TO_GRAYSCALE(diff)
    blur = GAUSSIAN_BLUR(gray, kernel=5x5)
    thresh = BINARY_THRESHOLD(blur, threshold=25)
    dilated = DILATE(thresh, iterations=3)
    contours = FIND_CONTOURS(dilated)
    motion_found = FALSE
    FOR each contour in contours:
        IF AREA(contour) >= MIN_AREA:
            DRAW_RECTANGLE(frame1, contour)
            motion_found = TRUE
    IF motion_found AND cooldown_elapsed:
        INCREMENT motion_count
        PLAY_ALERT_SOUND()
        SAVE_SCREENSHOT(frame1)
        START_RECORDING()
        LOG_EVENT("MOTION DETECTED")
    ELSE IF NOT motion_found AND recording:
        STOP_RECORDING()
    DISPLAY frame with dashboard overlay
    frame1 = frame2
    READ new frame2
RELEASE webcam
CLOSE all windows"""
p = doc.add_paragraph()
run = p.add_run(pseudo)
run.font.name = 'Courier New'
run.font.size = Pt(10)

# ============ CHAPTER 4: IMPLEMENTATION ============
page_break()
add_heading_custom("CHAPTER 4", 16)
add_heading_custom("IMPLEMENTATION", 16)
doc.add_paragraph()

add_heading_custom("4.1 Programming Language and Tools", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("Programming Language: Python 3.x", bold=True)
add_para("The following libraries and tools were used:")

table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
headers = ['Library/Tool', 'Version', 'Purpose']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
data = [
    ['OpenCV (cv2)', '4.13.0', 'Video capture, image processing, display'],
    ['Python', '3.x', 'Core programming language'],
    ['os', 'Built-in', 'File and directory management'],
    ['winsound', 'Built-in', 'Audio alert generation on Windows'],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

add_heading_custom("4.2 Code Modules and Functions", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()
table2 = doc.add_table(rows=6, cols=2)
table2.style = 'Table Grid'
table2.rows[0].cells[0].text = 'Function'
table2.rows[0].cells[1].text = 'Description'
table2.rows[0].cells[0].paragraphs[0].runs[0].bold = True
table2.rows[0].cells[1].paragraphs[0].runs[0].bold = True
funcs = [
    ['log_motion_event()', 'Logs events with timestamps to motion_log.txt'],
    ['start_recording()', 'Initializes VideoWriter and begins saving frames to .avi file'],
    ['stop_recording()', 'Releases VideoWriter and stops recording'],
    ['save_screenshot()', 'Saves current frame as a .jpg image with timestamp'],
    ['draw_dashboard()', 'Draws semi-transparent overlay with system info on each frame'],
]
for r, row_data in enumerate(funcs):
    for c, val in enumerate(row_data):
        table2.rows[r+1].cells[c].text = val

add_heading_custom("4.3 Key Code Snippets", 14, align=WD_ALIGN_PARAGRAPH.LEFT)

add_para("Motion Detection Core Logic:", bold=True)
code1 = """# Calculate difference between consecutive frames
diff = cv2.absdiff(frame1, frame2)
gray = cv2.cvtColor(diff, cv2.COLOR_BGR2GRAY)
blur = cv2.GaussianBlur(gray, (5, 5), 0)
_, thresh = cv2.threshold(blur, 25, 255, cv2.THRESH_BINARY)
dilated = cv2.dilate(thresh, None, iterations=3)
contours, _ = cv2.findContours(dilated, cv2.RETR_TREE,
                                cv2.CHAIN_APPROX_SIMPLE)"""
p = doc.add_paragraph()
run = p.add_run(code1)
run.font.name = 'Courier New'
run.font.size = Pt(10)

add_para("Contour Processing and Bounding Box:", bold=True)
code2 = """for contour in contours:
    area = cv2.contourArea(contour)
    if area < MIN_CONTOUR_AREA:
        continue
    (x, y, w, h) = cv2.boundingRect(contour)
    cv2.rectangle(frame1, (x, y), (x+w, y+h), (0, 255, 0), 2)"""
p = doc.add_paragraph()
run = p.add_run(code2)
run.font.name = 'Courier New'
run.font.size = Pt(10)

add_para("Dashboard Overlay:", bold=True)
code3 = """overlay = frame.copy()
cv2.rectangle(overlay, (0, 0), (frame_width, 80), (0,0,0), -1)
cv2.addWeighted(overlay, 0.6, frame, 0.4, 0, frame)
cv2.putText(frame, "SMART SURVEILLANCE SYSTEM", (10, 25),
            cv2.FONT_HERSHEY_SIMPLEX, 0.7, (0, 255, 255), 2)"""
p = doc.add_paragraph()
run = p.add_run(code3)
run.font.name = 'Courier New'
run.font.size = Pt(10)

# ============ CHAPTER 5: RESULTS AND ANALYSIS ============
page_break()
add_heading_custom("CHAPTER 5", 16)
add_heading_custom("RESULTS AND ANALYSIS", 16)
doc.add_paragraph()

add_heading_custom("5.1 Test Cases", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
table3 = doc.add_table(rows=6, cols=5)
table3.style = 'Table Grid'
tc_headers = ['Test Case No', 'Input', 'Expected Output', 'Actual Output', 'Result']
for i, h in enumerate(tc_headers):
    table3.rows[0].cells[i].text = h
    table3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
test_data = [
    ['1', 'No movement in frame', 'No motion detected, green status bar', 'Green "All Clear" bar displayed', 'Pass'],
    ['2', 'Person walks across frame', 'Motion detected with bounding box', 'Green rectangle drawn around person', 'Pass'],
    ['3', 'Motion detected', 'Screenshot saved automatically', 'Screenshot saved to screenshots/', 'Pass'],
    ['4', 'Motion detected', 'Video recording started', 'AVI file created in recordings/', 'Pass'],
    ['5', 'Press "q" key', 'System shuts down gracefully', 'System stopped, resources released', 'Pass'],
]
for r, row_data in enumerate(test_data):
    for c, val in enumerate(row_data):
        table3.rows[r+1].cells[c].text = val

add_heading_custom("5.2 Output Analysis", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("The system successfully demonstrated the following capabilities:")
add_bullet("Accurate detection of motion in the webcam feed with minimal false positives")
add_bullet("Real-time bounding box rendering around moving objects")
add_bullet("Automatic video recording triggered by motion events")
add_bullet("Screenshot capture with proper timestamp naming")
add_bullet("Comprehensive event logging with timestamps")
add_bullet("Smooth dashboard overlay without impacting video performance")
add_bullet("Clean system shutdown with proper resource release")

add_heading_custom("5.3 Complexity Analysis", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("Time Complexity:", bold=True)
table4 = doc.add_table(rows=5, cols=3)
table4.style = 'Table Grid'
cplx_headers = ['Operation', 'Best Case', 'Worst Case']
for i, h in enumerate(cplx_headers):
    table4.rows[0].cells[i].text = h
    table4.rows[0].cells[i].paragraphs[0].runs[0].bold = True
cplx_data = [
    ['Frame Differencing', 'O(N)', 'O(N)'],
    ['Gaussian Blur', 'O(N*k^2)', 'O(N*k^2)'],
    ['Contour Detection', 'O(N)', 'O(N)'],
    ['Overall per frame', 'O(N)', 'O(N)'],
]
for r, row_data in enumerate(cplx_data):
    for c, val in enumerate(row_data):
        table4.rows[r+1].cells[c].text = val
add_para("Where N = total number of pixels in the frame (width x height) and k = kernel size (5).")
add_para("Space Complexity: O(N) for storing frame buffers and intermediate images.")

# ============ CHAPTER 6: CONCLUSION ============
page_break()
add_heading_custom("CHAPTER 6", 16)
add_heading_custom("CONCLUSION AND FUTURE WORK", 16)
doc.add_paragraph()

add_heading_custom("6.1 Conclusion", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para("This project successfully demonstrates the development of a Smart Surveillance System using Python and OpenCV. The system effectively detects motion in real-time video feeds, automatically records video evidence, captures screenshots, generates audio alerts, and maintains detailed event logs. The frame differencing approach combined with image processing techniques provides reliable motion detection suitable for various surveillance scenarios.")
add_para("The system achieves its objectives of being cost-effective (using standard webcams), automated (no human monitoring required), and comprehensive (multiple output modes including video, screenshots, and logs). The dashboard overlay provides real-time system status information, making the system user-friendly and informative.")

add_heading_custom("6.2 Limitations", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
add_bullet("Sensitive to lighting changes (sudden changes in ambient light may trigger false positives)")
add_bullet("Requires a stationary camera (camera movement will be detected as motion)")
add_bullet("Cannot distinguish between different types of motion (human vs. animal vs. object)")
add_bullet("Performance depends on available system resources (CPU, RAM)")
add_bullet("Currently supports only single camera input")

add_heading_custom("6.3 Future Enhancements", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
add_bullet("Integration of deep learning models (YOLO, SSD) for object classification")
add_bullet("Face recognition for identifying known vs. unknown persons")
add_bullet("Email/SMS notification system for remote alerts")
add_bullet("Multi-camera support for comprehensive area coverage")
add_bullet("Cloud storage integration for remote access to recordings")
add_bullet("Night vision mode using infrared camera support")
add_bullet("Web-based dashboard for remote monitoring")
add_bullet("Adaptive thresholding to handle varying lighting conditions")

# ============ REFERENCES ============
page_break()
add_heading_custom("REFERENCES", 16)
doc.add_paragraph()

refs = [
    '1. G. Bradski and A. Kaehler, "Learning OpenCV: Computer Vision with the OpenCV Library," O\'Reilly Media, 2008.',
    '2. A. Sobral and A. Vacavant, "A comprehensive review of background subtraction algorithms evaluated with synthetic and real videos," Computer Vision and Image Understanding, Vol. 122, pp. 4-21, 2014.',
    '3. R. J. Radke, S. Andra, O. Al-Kofahi, and B. Roysam, "Image Change Detection Algorithms: A Systematic Survey," IEEE Transactions on Image Processing, Vol. 14, No. 3, pp. 294-307, 2005.',
    '4. OpenCV Documentation, https://docs.opencv.org/4.x/',
    '5. Python Documentation, https://docs.python.org/3/',
    '6. C. Stauffer and W. E. L. Grimson, "Adaptive background mixture models for real-time tracking," IEEE Computer Society Conference on Computer Vision and Pattern Recognition, 1999.',
]
for ref in refs:
    add_para(ref, align=WD_ALIGN_PARAGRAPH.LEFT)

# ============ SAVE ============
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Smart_Surveillance_System_Report.docx")
doc.save(output_path)
print(f"Report saved to: {output_path}")
